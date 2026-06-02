# src/chunking/pdf_page_chunker.py
#
# PDF 전용 page 기반 chunker입니다.
#
# 주요 목적:
# - PDF 문서를 페이지 단위로 안정적으로 청킹
# - section heading 과분할 문제 방지
# - page_start / page_end metadata 보존
# - 짧은 페이지는 다음 페이지와 병합
# - 긴 페이지는 max_chars 기준으로 재분할
# - embedding_text에 기관명/사업명/파일명/페이지 정보를 포함

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional
from html.parser import HTMLParser
import re
import shutil
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET

from src.utils.text_cleaner import clean_extracted_text


def _safe_str(value: Any) -> str:
    if value is None:
        return ""

    text = str(value)

    if text.lower() == "nan":
        return ""

    return text.strip()


def _normalize_text(text: Any) -> str:
    text = _safe_str(text)

    if not text:
        return ""

    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


_HANJA_REPLACEMENTS = {
    "第": "제",
    "條": "조",
    "項": "항",
    "號": "호",
    "章": "장",
    "節": "절",
    "別紙": "별지",
    "別添": "별첨",
    "以上": "이상",
    "以下": "이하",
    "未滿": "미만",
    "內": "내",
    "外": "외",
    "年": "년",
    "月": "월",
    "日": "일",
    "千": "천",
    "萬": "만",
    "億": "억",
    "一": "일",
    "二": "이",
    "三": "삼",
    "四": "사",
    "五": "오",
    "六": "육",
    "七": "칠",
    "八": "팔",
    "九": "구",
    "十": "십",
    "百": "백",
}

_KOREAN_DIGITS = {
    "영": 0,
    "공": 0,
    "일": 1,
    "이": 2,
    "삼": 3,
    "사": 4,
    "오": 5,
    "육": 6,
    "칠": 7,
    "팔": 8,
    "구": 9,
}
_KOREAN_SMALL_UNITS = {"십": 10, "백": 100, "천": 1000}
_KOREAN_LARGE_UNITS = {"만": 10_000, "억": 100_000_000, "조": 1_000_000_000_000}
_MONEY_UNITS = {
    "원": 1,
    "천원": 1_000,
    "만원": 10_000,
    "백만원": 1_000_000,
    "천만원": 10_000_000,
    "억원": 100_000_000,
}


def _parse_korean_number(text: str) -> int | None:
    text = re.sub(r"\s+", "", str(text or ""))
    total = 0
    section = 0
    number = 0
    consumed = False

    for char in text:
        if char in _KOREAN_DIGITS:
            number = _KOREAN_DIGITS[char]
            consumed = True
        elif char in _KOREAN_SMALL_UNITS:
            section += (number or 1) * _KOREAN_SMALL_UNITS[char]
            number = 0
            consumed = True
        elif char in _KOREAN_LARGE_UNITS:
            total += (section + number or 1) * _KOREAN_LARGE_UNITS[char]
            section = 0
            number = 0
            consumed = True
        else:
            return None

    if not consumed:
        return None

    return total + section + number


def _normalize_money_amounts(text: str) -> str:
    def numeric_repl(match: re.Match) -> str:
        number = int(match.group("number").replace(",", ""))
        return f"{number * _MONEY_UNITS[match.group('unit')]}원"

    def korean_repl(match: re.Match) -> str:
        number = _parse_korean_number(match.group("number"))

        if number is None:
            return match.group(0)

        return f"{number * _MONEY_UNITS[match.group('unit')]}원"

    text = re.sub(
        r"(?P<number>\d[\d,]*)\s*(?P<unit>억원|천만원|백만원|만원|천원|원)(?=\D|$)",
        numeric_repl,
        text,
    )
    text = re.sub(
        r"(?P<number>[영공일이삼사오육칠팔구십백천만억조]+)\s*(?P<unit>억원|천만원|백만원|만원|천원|원)(?=\D|$)",
        korean_repl,
        text,
    )
    return text


def _clean_markdown_text(text: Any) -> str:
    text = clean_extracted_text(str(text or ""))

    for source, target in sorted(_HANJA_REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(source, target)

    return _normalize_money_amounts(text)


def _extract_pdf_pages_with_pymupdf(pdf_path: str | Path) -> List[Dict[str, Any]]:
    """
    PyMuPDF를 사용해 PDF를 페이지별로 텍스트 추출합니다.

    Returns
    -------
    List[Dict[str, Any]]
        [
            {"page_no": 1, "text": "..."},
            {"page_no": 2, "text": "..."},
            ...
        ]
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ImportError(
            "pdf_page chunking을 사용하려면 PyMuPDF가 필요합니다. "
            "아래 명령으로 설치하세요:\n"
            "pip install pymupdf"
        ) from e

    pdf_path = Path(pdf_path)

    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 파일을 찾을 수 없습니다: {pdf_path}")

    pages: List[Dict[str, Any]] = []

    with fitz.open(pdf_path) as doc:
        for page_idx, page in enumerate(doc, start=1):
            # text 추출 모드:
            # - "text": 일반 텍스트 추출
            # - sort=True: 위에서 아래, 왼쪽에서 오른쪽 순서 정렬
            raw_text = page.get_text("text", sort=True) or ""
            clean_text = clean_extracted_text(raw_text)

            pages.append({
                "page_no": page_idx,
                "text": clean_text,
            })

    return pages


def _split_long_text_by_chars(
    text: str,
    max_chars: int,
    overlap_chars: int,
) -> List[str]:
    """
    긴 페이지 텍스트를 max_chars 기준으로 재분할합니다.
    문단 기준으로 먼저 묶고, 문단 하나가 너무 길면 문자 기준으로 강제 분할합니다.
    """
    text = _normalize_text(text)

    if not text:
        return []

    if len(text) <= max_chars:
        return [text]

    paragraphs = [
        paragraph.strip()
        for paragraph in re.split(r"\n\s*\n", text)
        if paragraph and paragraph.strip()
    ]

    if len(paragraphs) <= 1:
        paragraphs = [
            line.strip()
            for line in text.splitlines()
            if line and line.strip()
        ]

    chunks: List[str] = []
    current_parts: List[str] = []
    current_len = 0

    def flush_current() -> Optional[str]:
        if not current_parts:
            return None

        chunk_text = "\n\n".join(current_parts).strip()
        return chunk_text if chunk_text else None

    for paragraph in paragraphs:
        paragraph = paragraph.strip()

        if not paragraph:
            continue

        if len(paragraph) > max_chars:
            current_chunk = flush_current()

            if current_chunk:
                chunks.append(current_chunk)

            current_parts = []
            current_len = 0

            start = 0

            while start < len(paragraph):
                end = start + max_chars
                piece = paragraph[start:end].strip()

                if piece:
                    chunks.append(piece)

                if overlap_chars > 0:
                    next_start = end - overlap_chars
                else:
                    next_start = end

                if next_start <= start:
                    next_start = end

                start = next_start

            continue

        additional_len = len(paragraph) + 2

        if current_parts and current_len + additional_len > max_chars:
            current_chunk = flush_current()

            if current_chunk:
                chunks.append(current_chunk)

            overlap_text = ""

            if overlap_chars > 0 and current_chunk:
                overlap_text = current_chunk[-overlap_chars:].strip()

            current_parts = []
            current_len = 0

            if overlap_text:
                current_parts.append(overlap_text)
                current_len += len(overlap_text)

        current_parts.append(paragraph)
        current_len += additional_len

    last_chunk = flush_current()

    if last_chunk:
        chunks.append(last_chunk)

    return chunks


def _merge_short_pages(
    pages: List[Dict[str, Any]],
    min_chars: int,
    max_chars: int,
) -> List[Dict[str, Any]]:
    """
    너무 짧은 페이지 텍스트를 다음 페이지와 병합합니다.

    병합 기준:
    - 현재 buffer가 min_chars보다 짧으면 다음 페이지를 붙임
    - 단, 병합 후 max_chars를 크게 넘기면 병합하지 않음
    """
    merged_pages: List[Dict[str, Any]] = []

    buffer_text = ""
    buffer_start_page: Optional[int] = None
    buffer_end_page: Optional[int] = None

    def flush_buffer() -> None:
        nonlocal buffer_text, buffer_start_page, buffer_end_page

        text = _normalize_text(buffer_text)

        if text and buffer_start_page is not None and buffer_end_page is not None:
            merged_pages.append({
                "page_start": buffer_start_page,
                "page_end": buffer_end_page,
                "text": text,
            })

        buffer_text = ""
        buffer_start_page = None
        buffer_end_page = None

    for page in pages:
        page_no = int(page["page_no"])
        page_text = _normalize_text(page.get("text", ""))

        if not page_text:
            continue

        if buffer_start_page is None:
            buffer_text = page_text
            buffer_start_page = page_no
            buffer_end_page = page_no
            continue

        buffer_len = len(buffer_text)
        candidate_len = buffer_len + len(page_text) + 2

        should_merge = (
            buffer_len < min_chars
            and candidate_len <= max_chars
        )

        if should_merge:
            buffer_text = f"{buffer_text}\n\n{page_text}".strip()
            buffer_end_page = page_no
        else:
            flush_buffer()
            buffer_text = page_text
            buffer_start_page = page_no
            buffer_end_page = page_no

    flush_buffer()

    return merged_pages


def _build_embedding_text(
    text: str,
    organization: str = "",
    project_name: str = "",
    file_name: str = "",
    page_start: Optional[int] = None,
    page_end: Optional[int] = None,
    include_metadata: bool = True,
) -> str:
    text = _normalize_text(text)

    if not include_metadata:
        return text

    meta_lines = []

    if organization:
        meta_lines.append(f"기관명: {organization}")

    if project_name:
        meta_lines.append(f"사업명: {project_name}")

    if file_name:
        meta_lines.append(f"파일명: {file_name}")

    if page_start is not None and page_end is not None:
        if page_start == page_end:
            meta_lines.append(f"페이지: {page_start}")
        else:
            meta_lines.append(f"페이지: {page_start}-{page_end}")

    if not meta_lines:
        return text

    return "\n".join(meta_lines) + "\n\n본문:\n" + text


def chunk_pdf_by_page(
    doc_id: str,
    pdf_path: str | Path,
    file_name: str = "",
    file_type: str = "pdf",
    project_name: str = "",
    organization: str = "",
    max_chars: int = 3000,
    overlap_chars: int = 300,
    min_chars: int = 500,
    merge_short_pages: bool = True,
    include_metadata_in_embedding_text: bool = True,
) -> List[Dict[str, Any]]:
    """
    PDF 파일을 페이지 기반으로 chunking합니다.

    Parameters
    ----------
    doc_id:
        문서 ID입니다.

    pdf_path:
        PDF 파일 경로입니다.

    file_name:
        원본 파일명입니다.

    file_type:
        파일 형식입니다. 보통 pdf입니다.

    project_name:
        사업명입니다.

    organization:
        발주 기관명입니다.

    max_chars:
        chunk 최대 문자 수입니다. 페이지 텍스트가 이보다 길면 재분할합니다.

    overlap_chars:
        긴 페이지를 재분할할 때 overlap 문자 수입니다.

    min_chars:
        페이지 텍스트가 이보다 짧으면 다음 페이지와 병합합니다.

    merge_short_pages:
        짧은 페이지 병합 여부입니다.

    include_metadata_in_embedding_text:
        embedding_text에 metadata 포함 여부입니다.

    Returns
    -------
    List[Dict[str, Any]]
        chunk dict 목록입니다.
    """
    pdf_path = Path(pdf_path)

    pages = _extract_pdf_pages_with_pymupdf(pdf_path)

    if merge_short_pages:
        page_units = _merge_short_pages(
            pages=pages,
            min_chars=min_chars,
            max_chars=max_chars,
        )
    else:
        page_units = [
            {
                "page_start": int(page["page_no"]),
                "page_end": int(page["page_no"]),
                "text": _normalize_text(page.get("text", "")),
            }
            for page in pages
            if _normalize_text(page.get("text", ""))
        ]

    chunks: List[Dict[str, Any]] = []
    global_chunk_index = 0

    for page_unit in page_units:
        page_start = int(page_unit["page_start"])
        page_end = int(page_unit["page_end"])
        page_text = _normalize_text(page_unit.get("text", ""))

        if not page_text:
            continue

        split_texts = _split_long_text_by_chars(
            text=page_text,
            max_chars=max_chars,
            overlap_chars=overlap_chars,
        )

        for page_chunk_index, chunk_text in enumerate(split_texts):
            chunk_text = _normalize_text(chunk_text)

            if not chunk_text:
                continue

            if page_start == page_end:
                page_part = f"p{page_start:04d}"
            else:
                page_part = f"p{page_start:04d}_p{page_end:04d}"

            chunk_id = f"{doc_id}_{page_part}_c{page_chunk_index:03d}"

            embedding_text = _build_embedding_text(
                text=chunk_text,
                organization=organization,
                project_name=project_name,
                file_name=file_name,
                page_start=page_start,
                page_end=page_end,
                include_metadata=include_metadata_in_embedding_text,
            )

            chunks.append({
                "doc_id": str(doc_id),
                "chunk_id": chunk_id,
                "file_name": file_name,
                "file_type": file_type,
                "project_name": project_name,
                "organization": organization,
                "source_path": str(pdf_path),
                "page_start": page_start,
                "page_end": page_end,
                "page_chunk_index": page_chunk_index,
                "chunk_index": global_chunk_index,
                "section_id": f"page_{page_start:04d}_{page_end:04d}",
                "section_title": (
                    f"페이지 {page_start}"
                    if page_start == page_end
                    else f"페이지 {page_start}-{page_end}"
                ),
                "section_path": (
                    f"페이지 {page_start}"
                    if page_start == page_end
                    else f"페이지 {page_start}-{page_end}"
                ),
                "text": chunk_text,
                "embedding_text": embedding_text,
                "chunking_strategy": "pdf_page",
                "chunk_char_len": len(chunk_text),
            })

            global_chunk_index += 1

    return chunks


def analyze_pdf_page_split_lengths(
    pdf_path: str | Path,
    max_chars: int = 3000,
    overlap_chars: int = 300,
    min_chars: int = 500,
    merge_short_pages: bool = True,
) -> Dict[str, Any]:
    """
    PDF page 기반 chunking을 적용했을 때의 split 길이 통계를 계산합니다.

    최종 chunk 생성과 동일한 기준으로 page merge + long page split을 수행합니다.
    """
    pages = _extract_pdf_pages_with_pymupdf(pdf_path)

    if merge_short_pages:
        page_units = _merge_short_pages(
            pages=pages,
            min_chars=min_chars,
            max_chars=max_chars,
        )
    else:
        page_units = [
            {
                "page_start": int(page["page_no"]),
                "page_end": int(page["page_no"]),
                "text": _normalize_text(page.get("text", "")),
            }
            for page in pages
            if _normalize_text(page.get("text", ""))
        ]

    lengths: List[int] = []

    for page_unit in page_units:
        page_text = _normalize_text(page_unit.get("text", ""))

        split_texts = _split_long_text_by_chars(
            text=page_text,
            max_chars=max_chars,
            overlap_chars=overlap_chars,
        )

        for split_text in split_texts:
            split_text = _normalize_text(split_text)

            if split_text:
                lengths.append(len(split_text))

    if not lengths:
        return {
            "num_pages": len(pages),
            "pre_chunk_count": 0,
            "pre_chunk_min_chars": 0,
            "pre_chunk_max_chars": 0,
            "pre_chunk_mean_chars": 0.0,
            "pre_chunk_median_chars": 0.0,
            "pre_chunk_lengths": [],
        }

    sorted_lengths = sorted(lengths)
    n = len(sorted_lengths)

    if n % 2 == 1:
        median = float(sorted_lengths[n // 2])
    else:
        median = float((sorted_lengths[n // 2 - 1] + sorted_lengths[n // 2]) / 2)

    return {
        "num_pages": len(pages),
        "pre_chunk_count": len(lengths),
        "pre_chunk_min_chars": int(min(lengths)),
        "pre_chunk_max_chars": int(max(lengths)),
        "pre_chunk_mean_chars": float(sum(lengths) / len(lengths)),
        "pre_chunk_median_chars": median,
        "pre_chunk_lengths": lengths,
    }


# ---------------------------------------------------------
# HWP/PDF -> table-preserving markdown hierarchy chunker
# ---------------------------------------------------------

_ARTICLE_HEADING_RE = re.compile(r"^제\s*\d+\s*조(?:\s*\([^)]*\)|[^\n]*)?$")
_CIRCLED_ITEM_RE = re.compile(r"^[①-⑳]\s*")
_PAREN_NUMBER_ITEM_RE = re.compile(r"^\(?\d{1,2}\)\s+")
_NUMBER_ITEM_RE = re.compile(r"^\d{1,2}[.)]\s+")
_KOREAN_ITEM_RE = re.compile(r"^[가-힣][.)]\s+")
_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def _find_libreoffice_command() -> str:
    for command in ("libreoffice", "soffice"):
        resolved = shutil.which(command)

        if resolved:
            return resolved

    raise RuntimeError(
        "LibreOffice 실행 파일을 찾을 수 없습니다. "
        "가상 GPU 서버에 libreoffice/soffice가 PATH에 잡혀 있는지 확인하세요."
    )


def _convert_office_to_pdf_with_libreoffice(
    input_path: str | Path,
    output_dir: str | Path,
    timeout_sec: int = 180,
) -> Path:
    input_path = Path(input_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if not input_path.exists():
        raise FileNotFoundError(f"변환할 파일을 찾을 수 없습니다: {input_path}")

    command = _find_libreoffice_command()
    result = subprocess.run(
        [
            command,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(input_path),
        ],
        capture_output=True,
        text=True,
        timeout=timeout_sec,
        check=False,
    )

    converted_pdf = output_dir / f"{input_path.stem}.pdf"

    if result.returncode != 0 or not converted_pdf.exists():
        stdout = (result.stdout or "").strip()
        stderr = (result.stderr or "").strip()
        raise RuntimeError(
            "LibreOffice PDF 변환에 실패했습니다. "
            f"file={input_path}, stdout={stdout}, stderr={stderr}"
        )

    return converted_pdf


def _cell_to_markdown(value: Any) -> str:
    text = _clean_markdown_text(value)
    text = text.replace("\n", "<br>")
    text = text.replace("|", "\\|")
    return text


def _table_to_markdown(table: List[List[Any]]) -> str:
    rows = [
        [_cell_to_markdown(cell) for cell in row]
        for row in table
        if row and any(_safe_str(cell) for cell in row)
    ]

    if not rows:
        return ""

    max_cols = max(len(row) for row in rows)
    normalized_rows = [
        row + [""] * (max_cols - len(row))
        for row in rows
    ]

    header = normalized_rows[0]
    separator = ["---"] * max_cols
    body = normalized_rows[1:]

    markdown_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]

    for row in body:
        markdown_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(markdown_lines)


class _SimpleHTMLToMarkdownParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: List[str] = []
        self._text_parts: List[str] = []
        self._table: List[List[str]] | None = None
        self._row: List[str] | None = None
        self._cell_parts: List[str] | None = None

    def handle_starttag(self, tag: str, attrs: List[tuple[str, str | None]]) -> None:
        tag = tag.lower()

        if tag == "table":
            self._flush_text()
            self._table = []
        elif tag == "tr" and self._table is not None:
            self._row = []
        elif tag in {"td", "th"} and self._row is not None:
            self._cell_parts = []
        elif tag in {"p", "br", "div", "li", "h1", "h2", "h3"}:
            self._append_text("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()

        if tag in {"td", "th"} and self._row is not None and self._cell_parts is not None:
            self._row.append(_normalize_text(" ".join(self._cell_parts)))
            self._cell_parts = None
        elif tag == "tr" and self._table is not None and self._row is not None:
            if any(_safe_str(cell) for cell in self._row):
                self._table.append(self._row)
            self._row = None
        elif tag == "table" and self._table is not None:
            markdown_table = _table_to_markdown(self._table)
            if markdown_table:
                self.blocks.append(markdown_table)
            self._table = None
        elif tag in {"p", "div", "li", "h1", "h2", "h3"}:
            self._append_text("\n")

    def handle_data(self, data: str) -> None:
        if self._cell_parts is not None:
            self._cell_parts.append(data)
        else:
            self._append_text(data)

    def _append_text(self, text: str) -> None:
        if text:
            self._text_parts.append(text)

    def _flush_text(self) -> None:
        text = _clean_markdown_text("".join(self._text_parts))
        if text:
            self.blocks.append(text)
        self._text_parts = []

    def get_markdown(self) -> str:
        self._flush_text()
        return "\n\n".join(block for block in self.blocks if _normalize_text(block))


def _html_to_markdown(html_text: str) -> str:
    parser = _SimpleHTMLToMarkdownParser()
    parser.feed(str(html_text or ""))
    return parser.get_markdown()


def _run_command_text(command: List[str], timeout_sec: int = 180) -> str:
    try:
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout_sec,
            check=False,
        )
    except Exception:
        return ""

    if result.returncode != 0:
        return ""

    return result.stdout or ""


def _ocr_image_with_tesseract_cli(
    image_path: Path,
    language: str = "kor+eng",
    timeout_sec: int = 120,
) -> str:
    tesseract = shutil.which("tesseract")

    if not tesseract:
        return ""

    text = _run_command_text(
        [tesseract, str(image_path), "stdout", "-l", language],
        timeout_sec=timeout_sec,
    )
    return _clean_markdown_text(text)


def _extract_hwp_embedded_image_ocr(
    file_path: Path,
    language: str = "kor+eng",
    timeout_sec: int = 120,
) -> List[str]:
    try:
        import olefile
    except Exception:
        return []

    if not olefile.isOleFile(str(file_path)):
        return []

    image_texts: List[str] = []
    image_signatures = [
        (b"\x89PNG", ".png"),
        (b"\xff\xd8\xff", ".jpg"),
        (b"GIF8", ".gif"),
        (b"BM", ".bmp"),
    ]

    try:
        with olefile.OleFileIO(str(file_path)) as ole, tempfile.TemporaryDirectory(prefix="rfp_hwp_img_") as temp_dir:
            for stream_parts in ole.listdir():
                if not stream_parts or stream_parts[0] != "BinData":
                    continue

                data = ole.openstream("/".join(stream_parts)).read()
                suffix = ""

                for signature, candidate_suffix in image_signatures:
                    if data.startswith(signature):
                        suffix = candidate_suffix
                        break

                if not suffix:
                    continue

                image_path = Path(temp_dir) / f"{len(image_texts):04d}{suffix}"
                image_path.write_bytes(data)
                ocr_text = _ocr_image_with_tesseract_cli(
                    image_path=image_path,
                    language=language,
                    timeout_sec=timeout_sec,
                )

                if ocr_text:
                    image_texts.append(ocr_text)
    except Exception:
        return image_texts

    return image_texts


def _extract_hwp_markdown_with_external_tools(
    file_path: Path,
    timeout_sec: int = 180,
    ocr_language: str = "kor+eng",
) -> str:
    hwp5html = shutil.which("hwp5html")

    if hwp5html:
        with tempfile.TemporaryDirectory(prefix="rfp_hwp_html_") as temp_dir:
            result = subprocess.run(
                [hwp5html, "--output", temp_dir, str(file_path)],
                capture_output=True,
                text=True,
                timeout=timeout_sec,
                check=False,
            )

            if result.returncode == 0:
                html_parts = []
                for html_path in sorted(Path(temp_dir).rglob("*.html")):
                    html_parts.append(html_path.read_text(encoding="utf-8", errors="ignore"))

                image_ocr_parts = []
                for image_path in sorted(Path(temp_dir).rglob("*")):
                    if image_path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}:
                        continue

                    ocr_text = _ocr_image_with_tesseract_cli(
                        image_path=image_path,
                        language=ocr_language,
                        timeout_sec=timeout_sec,
                    )

                    if ocr_text:
                        image_ocr_parts.append(ocr_text)

                markdown = _html_to_markdown("\n".join(html_parts))
                markdown = "\n\n".join(part for part in [markdown, *image_ocr_parts] if part)
                if markdown:
                    return markdown

    hwp5txt = shutil.which("hwp5txt")

    if hwp5txt:
        text = _run_command_text([hwp5txt, str(file_path)], timeout_sec=timeout_sec)
        if text:
            return _clean_markdown_text(text)

    return ""


def _extract_hwp_markdown_with_ole(file_path: Path) -> str:
    try:
        from src.extractors.hwp_extractor import extract_hwp_text
    except Exception:
        return ""

    try:
        extracted = extract_hwp_text(file_path)
    except Exception:
        return ""

    text = _clean_markdown_text(extracted.get("text", ""))
    image_texts = _extract_hwp_embedded_image_ocr(file_path)
    return "\n\n".join(part for part in [text, *image_texts] if part)


def _extract_hwpx_pages_as_markdown(file_path: Path) -> List[Dict[str, Any]]:
    pages: List[Dict[str, Any]] = []

    try:
        with zipfile.ZipFile(file_path) as archive:
            xml_names = [
                name
                for name in archive.namelist()
                if name.lower().endswith(".xml") and (
                    "section" in name.lower() or "contents" in name.lower()
                )
            ]

            for page_idx, name in enumerate(sorted(xml_names), start=1):
                root = ET.fromstring(archive.read(name))
                paragraphs: List[str] = []
                tables: List[str] = []

                for elem in root.iter():
                    tag = elem.tag.split("}")[-1].lower()

                    if tag in {"p", "para"}:
                        text = _clean_markdown_text("".join(elem.itertext()))
                        if text:
                            paragraphs.append(text)
                    elif tag in {"tbl", "table"}:
                        rows: List[List[str]] = []
                        for row_elem in elem.iter():
                            row_tag = row_elem.tag.split("}")[-1].lower()
                            if row_tag not in {"tr", "row"}:
                                continue

                            row = []
                            for cell_elem in row_elem:
                                cell_tag = cell_elem.tag.split("}")[-1].lower()
                                if cell_tag in {"tc", "cell"}:
                                    row.append(_clean_markdown_text("".join(cell_elem.itertext())))

                            if any(row):
                                rows.append(row)

                        markdown_table = _table_to_markdown(rows)
                        if markdown_table:
                            tables.append(markdown_table)

                text = "\n\n".join([*paragraphs, *tables]).strip()
                if text:
                    pages.append({"page_no": page_idx, "text": text, "table_count": len(tables)})
    except Exception:
        return []

    return pages


def _extract_docx_pages_as_markdown(file_path: Path) -> List[Dict[str, Any]]:
    try:
        from docx import Document
    except Exception:
        return []

    try:
        document = Document(str(file_path))
    except Exception:
        return []

    parts = [
        _clean_markdown_text(paragraph.text)
        for paragraph in document.paragraphs
        if _clean_markdown_text(paragraph.text)
    ]
    tables = []

    for table in document.tables:
        rows = [
            [_clean_markdown_text(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        markdown_table = _table_to_markdown(rows)
        if markdown_table:
            tables.append(markdown_table)

    text = "\n\n".join([*parts, *tables]).strip()

    return [{"page_no": 1, "text": text, "table_count": len(tables)}] if text else []


def _extract_text_or_html_pages_as_markdown(file_path: Path) -> List[Dict[str, Any]]:
    text = file_path.read_text(encoding="utf-8", errors="ignore")

    if file_path.suffix.lower() in {".html", ".htm"}:
        text = _html_to_markdown(text)
    else:
        text = _clean_markdown_text(text)

    return [{"page_no": 1, "text": text, "table_count": text.count("| ---")}] if text else []


def _extract_hwp_pages_as_markdown_without_libreoffice(
    file_path: Path,
    timeout_sec: int = 180,
    ocr_language: str = "kor+eng",
) -> List[Dict[str, Any]]:
    markdown = _extract_hwp_markdown_with_external_tools(
        file_path=file_path,
        timeout_sec=timeout_sec,
        ocr_language=ocr_language,
    )

    if not markdown:
        markdown = _extract_hwp_markdown_with_ole(file_path)

    if not markdown:
        return []

    return [{"page_no": 1, "text": markdown, "table_count": markdown.count("| ---")}]


def _extract_pdf_pages_as_markdown_with_pdfplumber(
    pdf_path: str | Path,
    enable_ocr_fallback: bool = True,
    ocr_min_text_chars: int = 40,
    ocr_dpi: int = 300,
    ocr_language: str = "kor+eng",
) -> List[Dict[str, Any]]:
    try:
        import pdfplumber
    except ImportError as e:
        raise ImportError(
            "표 복원형 PDF/HWP 청킹에는 pdfplumber가 필요합니다. "
            "requirements.txt 설치 상태를 확인하세요."
        ) from e

    pdf_path = Path(pdf_path)

    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 파일을 찾을 수 없습니다: {pdf_path}")

    pages: List[Dict[str, Any]] = []

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            text = _clean_markdown_text(page.extract_text() or "")

            table_markdowns: List[str] = []

            for table_idx, table in enumerate(page.extract_tables() or [], start=1):
                markdown_table = _table_to_markdown(table)

                if markdown_table:
                    table_markdowns.append(f"표 {table_idx}\n{markdown_table}")

            ocr_used = False

            if (
                enable_ocr_fallback
                and len(text) < ocr_min_text_chars
                and not table_markdowns
            ):
                ocr_text = _extract_pdf_page_text_with_pymupdf_ocr(
                    pdf_path=pdf_path,
                    page_no=page_idx,
                    dpi=ocr_dpi,
                    language=ocr_language,
                )

                if len(ocr_text) > len(text):
                    text = ocr_text
                    ocr_used = True

            parts = [part for part in [text, *table_markdowns] if _normalize_text(part)]

            pages.append({
                "page_no": page_idx,
                "text": "\n\n".join(parts).strip(),
                "table_count": len(table_markdowns),
                "ocr_used": ocr_used,
            })

    return pages


def _extract_pdf_page_text_with_pymupdf_ocr(
    pdf_path: str | Path,
    page_no: int,
    dpi: int = 300,
    language: str = "kor+eng",
) -> str:
    """
    Optional Hi-Res OCR fallback.

    No Python OCR dependency is required. This uses PyMuPDF's OCR bridge only
    when the installed PyMuPDF supports it and Tesseract is available on the
    machine.
    """
    try:
        import fitz
    except ImportError:
        return ""

    try:
        with fitz.open(pdf_path) as doc:
            if page_no < 1 or page_no > len(doc):
                return ""

            page = doc[page_no - 1]

            if not hasattr(page, "get_textpage_ocr"):
                return ""

            textpage = page.get_textpage_ocr(
                flags=0,
                language=language,
                dpi=int(dpi),
                full=True,
            )
            text = page.get_text("text", textpage=textpage, sort=True) or ""
            return _clean_markdown_text(text)
    except Exception:
        return ""


def _looks_like_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _classify_clause_heading(line: str) -> tuple[int, str] | None:
    stripped = line.strip()

    if not stripped:
        return None

    if _ARTICLE_HEADING_RE.match(stripped):
        return 1, stripped

    if _CIRCLED_ITEM_RE.match(stripped) or _PAREN_NUMBER_ITEM_RE.match(stripped):
        return 2, stripped

    if _NUMBER_ITEM_RE.match(stripped) or _KOREAN_ITEM_RE.match(stripped):
        return 3, stripped

    return None


def _pages_to_clause_markdown_lines(
    pages: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    markdown_lines: List[Dict[str, Any]] = []

    for page in pages:
        page_no = int(page["page_no"])
        text = _normalize_text(page.get("text", ""))

        if not text:
            continue

        for raw_line in text.splitlines():
            line = raw_line.strip()

            if not line:
                markdown_lines.append({"page_no": page_no, "line": ""})
                continue

            if line.startswith("#") or _looks_like_markdown_table_line(line):
                markdown_lines.append({"page_no": page_no, "line": line})
                continue

            classified = _classify_clause_heading(line)

            if classified is None:
                markdown_lines.append({"page_no": page_no, "line": line})
                continue

            level, title = classified
            markdown_lines.append({
                "page_no": page_no,
                "line": f"{'#' * level} {title}",
            })

    return markdown_lines


def _split_markdown_lines_into_sections(
    markdown_lines: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    current_lines: List[str] = []
    current_page_start: Optional[int] = None
    current_page_end: Optional[int] = None
    current_title = "문서 시작"
    current_level = 0
    heading_stack: List[tuple[int, str]] = []

    def current_path() -> List[str]:
        return [title for _, title in heading_stack]

    def flush() -> None:
        nonlocal current_lines, current_page_start, current_page_end

        text = _normalize_text("\n".join(current_lines))

        if text and current_page_start is not None and current_page_end is not None:
            sections.append({
                "section_title": current_title,
                "section_level": current_level,
                "section_path": current_path() or [current_title],
                "page_start": current_page_start,
                "page_end": current_page_end,
                "text": text,
            })

        current_lines = []
        current_page_start = None
        current_page_end = None

    for item in markdown_lines:
        line = item["line"]
        page_no = int(item["page_no"])
        heading_match = _MARKDOWN_HEADING_RE.match(line)

        if heading_match:
            flush()

            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            heading_stack = [
                stack_item
                for stack_item in heading_stack
                if stack_item[0] < level
            ]
            heading_stack.append((level, title))

            current_title = title
            current_level = level

        if current_page_start is None:
            current_page_start = page_no

        current_page_end = page_no
        current_lines.append(line)

    flush()

    return sections


def _split_text_table_aware(
    text: str,
    max_chars: int,
    overlap_chars: int,
) -> List[str]:
    text = _normalize_text(text)

    if not text:
        return []

    if len(text) <= max_chars:
        return [text]

    blocks: List[str] = []
    current_table: List[str] = []
    current_text: List[str] = []

    def flush_text() -> None:
        nonlocal current_text
        block = _normalize_text("\n".join(current_text))

        if block:
            blocks.append(block)

        current_text = []

    def flush_table() -> None:
        nonlocal current_table
        block = _normalize_text("\n".join(current_table))

        if block:
            blocks.append(block)

        current_table = []

    for line in text.splitlines():
        if _looks_like_markdown_table_line(line):
            flush_text()
            current_table.append(line)
        else:
            flush_table()
            current_text.append(line)

    flush_text()
    flush_table()

    chunks: List[str] = []
    current_parts: List[str] = []

    def flush_current() -> None:
        nonlocal current_parts
        chunk = _normalize_text("\n\n".join(current_parts))

        if chunk:
            chunks.append(chunk)

        current_parts = []

    for block in blocks:
        if len(block) > max_chars:
            flush_current()
            chunks.extend(_split_long_text_by_chars(block, max_chars, overlap_chars))
            continue

        candidate = _normalize_text("\n\n".join([*current_parts, block]))

        if current_parts and len(candidate) > max_chars:
            flush_current()

        current_parts.append(block)

    flush_current()

    return chunks


def _select_parent_context(
    parent_text: str,
    child_text: str,
    max_chars: int,
) -> tuple[str, bool]:
    parent_text = _normalize_text(parent_text)
    child_text = _normalize_text(child_text)

    if max_chars <= 0 or len(parent_text) <= max_chars:
        return parent_text, False

    child_pos = parent_text.find(child_text[:80]) if child_text else -1

    if child_pos < 0:
        return parent_text[:max_chars].strip(), True

    half = max_chars // 2
    start = max(0, child_pos - half)
    end = min(len(parent_text), start + max_chars)
    start = max(0, end - max_chars)
    context = parent_text[start:end].strip()

    if start > 0:
        context = "... " + context

    if end < len(parent_text):
        context = context + " ..."

    return context, True


def _build_markdown_hierarchy_chunks(
    doc_id: str,
    sections: List[Dict[str, Any]],
    source_path: str | Path,
    file_name: str = "",
    file_type: str = "",
    project_name: str = "",
    organization: str = "",
    max_chars: int = 1500,
    overlap_chars: int = 150,
    min_chars: int = 30,
    parent_child_threshold: int = 2000,
    parent_child_enabled: bool = True,
    child_max_chars: int = 300,
    max_parent_context_chars: int = 1800,
    store_parent_context_in_metadata: bool = False,
    store_child_text: bool = False,
    include_metadata_in_embedding_text: bool = True,
    converted_pdf_path: str | Path | None = None,
) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    global_chunk_index = 0
    source_path = Path(source_path)

    for section_idx, section in enumerate(sections):
        section_text = _clean_markdown_text(section.get("text", ""))

        if len(section_text) < min_chars:
            continue

        is_parent_child = parent_child_enabled and len(section_text) > parent_child_threshold
        effective_max_chars = child_max_chars if is_parent_child else max_chars
        split_texts = _split_text_table_aware(
            text=section_text,
            max_chars=effective_max_chars,
            overlap_chars=overlap_chars,
        )

        for split_idx, chunk_text in enumerate(split_texts):
            chunk_text = _clean_markdown_text(chunk_text)

            if len(chunk_text) < min_chars:
                continue

            page_start = int(section["page_start"])
            page_end = int(section["page_end"])
            section_path = section.get("section_path") or [section.get("section_title", "")]
            section_path_text = " > ".join(str(x) for x in section_path if x)
            chunk_id = (
                f"{doc_id}_mdsec_{section_idx:04d}_"
                f"p{page_start:04d}_p{page_end:04d}_c{split_idx:02d}"
            )

            embedding_text = _build_embedding_text(
                text=chunk_text,
                organization=organization,
                project_name=project_name,
                file_name=file_name,
                page_start=page_start,
                page_end=page_end,
                include_metadata=include_metadata_in_embedding_text,
            )

            if include_metadata_in_embedding_text and section_path_text:
                embedding_text = (
                    embedding_text.replace(
                        "\n\n본문:\n",
                        f"\n섹션경로: {section_path_text}\n\n본문:\n",
                        1,
                    )
                    if "\n\n본문:\n" in embedding_text
                    else f"섹션경로: {section_path_text}\n\n{embedding_text}"
                )

            metadata = {
                "source_path": str(source_path),
                "converted_pdf_path": str(converted_pdf_path or ""),
                "page_start": page_start,
                "page_end": page_end,
                "section_path": section_path,
                "section_title": section.get("section_title", ""),
                "section_level": section.get("section_level", 0),
                "chunk_strategy": (
                    "markdown_parent_child_table_aware"
                    if is_parent_child
                    else "markdown_hierarchical"
                ),
                "has_table": "|" in chunk_text and "---" in chunk_text,
            }

            if is_parent_child:
                llm_text, parent_truncated = _select_parent_context(
                    parent_text=section_text,
                    child_text=chunk_text,
                    max_chars=max_parent_context_chars,
                )
                metadata["parent_context_char_len"] = len(section_text)
                metadata["parent_context_truncated"] = parent_truncated

                if store_parent_context_in_metadata:
                    metadata["parent_context"] = llm_text
            else:
                llm_text = chunk_text

            chunks.append({
                "doc_id": str(doc_id),
                "chunk_id": chunk_id,
                "file_name": file_name,
                "file_type": file_type,
                "project_name": project_name,
                "organization": organization,
                "source_path": str(source_path),
                "converted_pdf_path": str(converted_pdf_path or ""),
                "page_start": page_start,
                "page_end": page_end,
                "page_chunk_index": split_idx,
                "chunk_index": global_chunk_index,
                "section_id": f"{doc_id}_mdsec_{section_idx:04d}",
                "section_title": section.get("section_title", ""),
                "section_path": section_path,
                "section_path_with_level": [
                    {"level": idx + 1, "title": title}
                    for idx, title in enumerate(section_path)
                ],
                "section_level": section.get("section_level", 0),
                "split_idx": split_idx,
                "text": llm_text,
                "child_text": chunk_text if is_parent_child and store_child_text else "",
                "parent_text": "",
                "embedding_text": embedding_text,
                "metadata": metadata,
                "chunking_method": "hwp_pdf_markdown_hierarchy",
                "chunking_strategy": (
                    "hwp_pdf_markdown_parent_child"
                    if is_parent_child
                    else "hwp_pdf_markdown_hierarchy"
                ),
                "chunk_char_len": len(chunk_text),
                "char_count": len(llm_text),
                "char_len": len(llm_text),
            })

            global_chunk_index += 1

    return chunks


def chunk_hwp_or_pdf_by_markdown_hierarchy(
    doc_id: str,
    file_path: str | Path,
    file_name: str = "",
    file_type: str = "",
    project_name: str = "",
    organization: str = "",
    max_chars: int = 1500,
    overlap_chars: int = 150,
    min_chars: int = 30,
    parent_child_threshold: int = 2000,
    parent_child_enabled: bool = True,
    child_max_chars: int = 300,
    max_parent_context_chars: int = 1800,
    store_parent_context_in_metadata: bool = False,
    store_child_text: bool = False,
    conversion_dir: str | Path | None = None,
    keep_converted_pdf: bool = False,
    libreoffice_timeout_sec: int = 180,
    include_metadata_in_embedding_text: bool = True,
    enable_ocr_fallback: bool = True,
    ocr_min_text_chars: int = 40,
    ocr_dpi: int = 300,
    ocr_language: str = "kor+eng",
) -> List[Dict[str, Any]]:
    """
    HWP/PDF를 표 보존형 PDF 파싱 결과로 정규화한 뒤 마크다운 계층 구조로 청킹합니다.

    PDF는 pdfplumber 기반으로 처리하고, HWP/HWPX/DOCX/TXT/HTML은 LibreOffice 없이
    텍스트/HTML/XML 기반 markdown으로 바로 처리합니다.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    file_type = file_type or suffix.lstrip(".")

    temp_dir_obj: tempfile.TemporaryDirectory[str] | None = None
    converted_pdf_path: Path | None = None

    try:
        if suffix == ".pdf":
            pages = _extract_pdf_pages_as_markdown_with_pdfplumber(
                file_path,
                enable_ocr_fallback=enable_ocr_fallback,
                ocr_min_text_chars=ocr_min_text_chars,
                ocr_dpi=ocr_dpi,
                ocr_language=ocr_language,
            )
            converted_pdf_path = file_path
        elif suffix == ".hwp":
            pages = _extract_hwp_pages_as_markdown_without_libreoffice(
                file_path=file_path,
                timeout_sec=libreoffice_timeout_sec,
                ocr_language=ocr_language,
            )
        elif suffix == ".hwpx":
            pages = _extract_hwpx_pages_as_markdown(file_path)
        elif suffix == ".docx":
            pages = _extract_docx_pages_as_markdown(file_path)
        elif suffix in {".txt", ".md", ".html", ".htm"}:
            pages = _extract_text_or_html_pages_as_markdown(file_path)
        else:
            raise ValueError(
                "markdown hierarchy 전략은 pdf/hwp/hwpx/docx/txt/html 파일만 지원합니다. "
                f"file_path={file_path}"
            )

        if not pages:
            raise ValueError(f"markdown으로 추출된 본문이 없습니다: {file_path}")

        markdown_lines = _pages_to_clause_markdown_lines(pages)
        sections = _split_markdown_lines_into_sections(markdown_lines)

        return _build_markdown_hierarchy_chunks(
            doc_id=doc_id,
            sections=sections,
            source_path=file_path,
            file_name=file_name,
            file_type=file_type,
            project_name=project_name,
            organization=organization,
            max_chars=max_chars,
            overlap_chars=overlap_chars,
            min_chars=min_chars,
            parent_child_threshold=parent_child_threshold,
            parent_child_enabled=parent_child_enabled,
            child_max_chars=child_max_chars,
            max_parent_context_chars=max_parent_context_chars,
            store_parent_context_in_metadata=store_parent_context_in_metadata,
            store_child_text=store_child_text,
            include_metadata_in_embedding_text=include_metadata_in_embedding_text,
            converted_pdf_path=converted_pdf_path,
        )
    finally:
        if temp_dir_obj is not None and not keep_converted_pdf:
            temp_dir_obj.cleanup()


def analyze_markdown_hierarchy_split_lengths(
    file_path: str | Path,
    max_chars: int = 1500,
    overlap_chars: int = 150,
    min_chars: int = 30,
    parent_child_threshold: int = 2000,
    parent_child_enabled: bool = True,
    child_max_chars: int = 300,
    max_parent_context_chars: int = 1800,
    store_parent_context_in_metadata: bool = False,
    store_child_text: bool = False,
    conversion_dir: str | Path | None = None,
    keep_converted_pdf: bool = False,
    libreoffice_timeout_sec: int = 180,
) -> Dict[str, Any]:
    chunks = chunk_hwp_or_pdf_by_markdown_hierarchy(
        doc_id="analysis",
        file_path=file_path,
        max_chars=max_chars,
        overlap_chars=overlap_chars,
        min_chars=min_chars,
        parent_child_threshold=parent_child_threshold,
        parent_child_enabled=parent_child_enabled,
        child_max_chars=child_max_chars,
        max_parent_context_chars=max_parent_context_chars,
        store_parent_context_in_metadata=store_parent_context_in_metadata,
        store_child_text=store_child_text,
        conversion_dir=conversion_dir,
        keep_converted_pdf=keep_converted_pdf,
        libreoffice_timeout_sec=libreoffice_timeout_sec,
        include_metadata_in_embedding_text=False,
    )

    lengths = [
        len(_normalize_text(chunk.get("text", "")))
        for chunk in chunks
        if _normalize_text(chunk.get("text", ""))
    ]
    pages = {
        int(page_no)
        for chunk in chunks
        for page_no in range(
            int(chunk.get("page_start") or 0),
            int(chunk.get("page_end") or 0) + 1,
        )
        if page_no > 0
    }
    sections = {
        chunk.get("section_id")
        for chunk in chunks
        if chunk.get("section_id")
    }

    if not lengths:
        return {
            "num_pages": len(pages),
            "num_sections": len(sections),
            "pre_chunk_count": 0,
            "pre_chunk_min_chars": 0,
            "pre_chunk_max_chars": 0,
            "pre_chunk_mean_chars": 0.0,
            "pre_chunk_median_chars": 0.0,
            "pre_chunk_lengths": [],
        }

    sorted_lengths = sorted(lengths)
    n = len(sorted_lengths)
    median = (
        float(sorted_lengths[n // 2])
        if n % 2 == 1
        else float((sorted_lengths[n // 2 - 1] + sorted_lengths[n // 2]) / 2)
    )

    return {
        "num_pages": len(pages),
        "num_sections": len(sections),
        "pre_chunk_count": len(lengths),
        "pre_chunk_min_chars": int(min(lengths)),
        "pre_chunk_max_chars": int(max(lengths)),
        "pre_chunk_mean_chars": float(sum(lengths) / len(lengths)),
        "pre_chunk_median_chars": median,
        "pre_chunk_lengths": lengths,
    }
